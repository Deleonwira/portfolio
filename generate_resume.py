import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def create_element(name):
    return OxmlElement(name)

def set_cell_border(cell, **kwargs):
    """
    kwargs: top, bottom, left, right
    values: dict(sz=12, val='single', color='FF0000', space='0')
    """
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key, attr in [('sz', 'w:sz'), ('val', 'w:val'), ('color', 'w:color'), ('space', 'w:space')]:
                if key in edge_data:
                    element.set(qn(attr), str(edge_data[key]))

def add_bottom_border(paragraph, color_hex="1A365D", size_pt=1.5):
    """Adds a bottom border to a paragraph for elegant section headers."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(int(size_pt * 8)))  # Eighths of a point
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)

def build_resume():
    doc = Document()

    # Page Margins (0.6 inch margins for modern clean look)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)

    # Styling colors
    PRIMARY_COLOR = RGBColor(26, 54, 93)      # Deep Navy/Slate (#1A365D)
    SECONDARY_COLOR = RGBColor(43, 108, 176)  # Blue Accent (#2B6CB0)
    DARK_TEXT = RGBColor(45, 55, 72)         # Off-black Charcoal (#2D3748)
    MUTED_TEXT = RGBColor(113, 128, 150)     # Cool Gray (#718096)

    FONT_FAMILY = 'Arial'

    # Normal Style configuration
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_FAMILY
    font.size = Pt(10)
    font.color.rgb = DARK_TEXT

    # -------------------------------------------------------------
    # HEADER SECTION
    # -------------------------------------------------------------
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(2)
    title_run = title_p.add_run("HAMZA DELEON WIRADARMA")
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = PRIMARY_COLOR

    subtitle_p = doc.add_paragraph()
    subtitle_p.paragraph_format.space_before = Pt(0)
    subtitle_p.paragraph_format.space_after = Pt(6)
    sub_run = subtitle_p.add_run("Full Stack Developer")
    sub_run.font.size = Pt(13)
    sub_run.font.bold = True
    sub_run.font.color.rgb = SECONDARY_COLOR

    contact_p = doc.add_paragraph()
    contact_p.paragraph_format.space_before = Pt(0)
    contact_p.paragraph_format.space_after = Pt(12)
    
    info_parts = [
        ("📍 Tangerang, Banten, 15116", False),
        ("  |  ", True),
        ("📞 +62 831-9463-4063", False),
        ("  |  ", True),
        ("🌐 Indonesian", False),
        ("  |  ", True),
        ("🎂 Feb 23, 2006", False)
    ]
    for text, is_muted in info_parts:
        r = contact_p.add_run(text)
        r.font.size = Pt(9.5)
        if is_muted:
            r.font.color.rgb = MUTED_TEXT
        else:
            r.font.color.rgb = DARK_TEXT

    # -------------------------------------------------------------
    # HELPER FOR SECTION HEADINGS
    # -------------------------------------------------------------
    def add_section_heading(heading_text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(4)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(heading_text.upper())
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = PRIMARY_COLOR
        add_bottom_border(h, color_hex="2B6CB0", size_pt=1.2)
        return h

    # -------------------------------------------------------------
    # PROFESSIONAL SUMMARY
    # -------------------------------------------------------------
    add_section_heading("Professional Summary")
    
    summary_p = doc.add_paragraph()
    summary_p.paragraph_format.space_before = Pt(4)
    summary_p.paragraph_format.space_after = Pt(8)
    summary_p.paragraph_format.line_spacing = 1.15
    s_run = summary_p.add_run(
        "Aspiring Full Stack Developer with a strong technical foundation in React.js, Express.js, Node.js, Python, PHP (Laravel), Java (Spring Boot), and relational/NoSQL databases. Experienced in building full-stack web applications, REST APIs, interactive client-side systems, and decision-support algorithms through academic and freelance projects. Passionate about solving real-world challenges, continuously acquiring modern technologies, and writing clean, maintainable code. Seeking opportunities to contribute to high-impact software engineering teams while expanding expertise in full-stack architecture and AI."
    )
    s_run.font.size = Pt(9.5)

    # -------------------------------------------------------------
    # TECHNICAL SKILLS
    # -------------------------------------------------------------
    add_section_heading("Technical Skills")

    skills_data = [
        ("Frontend Development", "React.js (★★★★☆), JavaScript (ES6+), HTML5, CSS3 (Flexbox/Grid), TailwindCSS, Bootstrap"),
        ("Backend & Languages", "Express.js (★★★★☆), Python (★★★★☆), PHP / Laravel (★★★☆☆), Java / Spring Boot (★★★☆☆), Visual Basic (★★★☆☆)"),
        ("Databases & Storage", "MySQL (★★★★☆), MongoDB, IndexedDB"),
        ("Developer Tools", "Git & GitHub (★★★★☆), REST APIs, Jest, React Testing Library, Postman, Vite")
    ]

    for category, details in skills_data:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.1)
        
        c_run = p.add_run(f"•  {category}: ")
        c_run.font.bold = True
        c_run.font.size = Pt(9.5)
        c_run.font.color.rgb = PRIMARY_COLOR

        d_run = p.add_run(details)
        d_run.font.size = Pt(9.5)

    # -------------------------------------------------------------
    # WORK EXPERIENCE
    # -------------------------------------------------------------
    add_section_heading("Work Experience")

    # --- Role 1: Freelance ---
    role1_h = doc.add_paragraph()
    role1_h.paragraph_format.space_before = Pt(6)
    role1_h.paragraph_format.space_after = Pt(1)
    role1_h.paragraph_format.keep_with_next = True
    
    r1_title = role1_h.add_run("Fullstack Developer — Freelance")
    r1_title.font.bold = True
    r1_title.font.size = Pt(10.5)
    r1_title.font.color.rgb = PRIMARY_COLOR

    r1_loc = role1_h.add_run(" | Tangerang, Banten")
    r1_loc.font.size = Pt(10)
    r1_loc.font.color.rgb = MUTED_TEXT

    # Date float right / separate run
    r1_date = role1_h.add_run(" \t2023 – Present")
    r1_date.font.bold = True
    r1_date.font.size = Pt(9.5)
    r1_date.font.color.rgb = SECONDARY_COLOR

    # Freelance projects bullets
    projects = [
        (
            "Photon — Digital Image Processing Studio",
            "Designed and built a browser-based, desktop-class image processing platform combining a pixel-level HTML5 Canvas engine with a Python/Flask/OpenCV backend. Implemented 6 from-scratch edge detection algorithms, 5 custom binary image compression codecs (Huffman, LZW, RLE), YOLOv4-tiny AI object recognition, real-time RGB histogram analysis, and a framework-free Pub/Sub state management system supporting IndexedDB and MySQL storage."
        ),
        (
            "Decision Support System for Bandsaw Machine Procurement (MOORA Algorithm)",
            "Developed a full-stack web decision support application utilizing the MOORA mathematical decision algorithm to optimize industrial machinery purchasing. Built with PHP and MySQL, the system normalizes multi-attribute cost/benefit criteria (price, cutting capacity, power efficiency, durability) to generate objective, data-driven machine procurement rankings."
        ),
        (
            "Ogani — Full-Stack Organic E-Commerce Platform",
            "Built a complete full-stack e-commerce application specializing in organic grocery retail. Engineered customer-facing features including product catalog filtering, dynamic shopping cart, multi-step checkout, and order history tracking, paired with a comprehensive admin dashboard for real-time inventory management, order processing, and user authorization."
        )
    ]

    for p_title, p_desc in projects:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_before = Pt(2)
        bp.paragraph_format.space_after = Pt(3)
        bp.paragraph_format.left_indent = Inches(0.25)
        
        t_run = bp.add_run(f"{p_title}: ")
        t_run.font.bold = True
        t_run.font.size = Pt(9.5)
        t_run.font.color.rgb = DARK_TEXT

        desc_run = bp.add_run(p_desc)
        desc_run.font.size = Pt(9.5)

    # --- Role 2: Frontend Developer Intern ---
    role2_h = doc.add_paragraph()
    role2_h.paragraph_format.space_before = Pt(8)
    role2_h.paragraph_format.space_after = Pt(1)
    role2_h.paragraph_format.keep_with_next = True
    
    r2_title = role2_h.add_run("Frontend Developer Intern")
    r2_title.font.bold = True
    r2_title.font.size = Pt(10.5)
    r2_title.font.color.rgb = PRIMARY_COLOR

    r2_loc = role2_h.add_run(" | Meruya, Jakarta Barat")
    r2_loc.font.size = Pt(10)
    r2_loc.font.color.rgb = MUTED_TEXT

    r2_date = role2_h.add_run(" \tJuly 2021 – November 2021")
    r2_date.font.bold = True
    r2_date.font.size = Pt(9.5)
    r2_date.font.color.rgb = SECONDARY_COLOR

    intern_bullets = [
        "Developed responsive, accessible UI components using HTML5, CSS3 (Flexbox/Grid), and modern JavaScript (ES6+); implemented reusable React components that accelerated feature delivery by ~20%.",
        "Integrated front-end interfaces with RESTful APIs and implemented robust client-side validation and error handling, enhancing user task completion and reducing support tickets.",
        "Optimized front-end performance via lazy loading, code-splitting, and asset minification, reducing initial page load time by ~25%.",
        "Wrote unit and integration tests (Jest, React Testing Library) and maintained component documentation to improve codebase maintainability and team onboarding.",
        "Resolved cross-browser compatibility issues and addressed accessibility (WCAG compliance) improvements prior to software releases."
    ]

    for b in intern_bullets:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_before = Pt(1)
        bp.paragraph_format.space_after = Pt(2)
        bp.paragraph_format.left_indent = Inches(0.25)
        
        b_run = bp.add_run(b)
        b_run.font.size = Pt(9.5)

    # -------------------------------------------------------------
    # EDUCATION
    # -------------------------------------------------------------
    add_section_heading("Education")

    edu_items = [
        (
            "JAKARTA STATE POLYTECHNIC (Politeknik Negeri Jakarta - PNJ)",
            "Depok, Indonesia",
            "Bachelor of Applied Science in Informatics Engineering (Teknik Informatika)",
            "July 2022 – Present (Active)"
        ),
        (
            "CCIT — FACULTY OF ENGINEERING, UNIVERSITY OF INDONESIA (UI)",
            "Depok, Indonesia",
            "Professional Program in Information Technology",
            "July 2022 – July 2026"
        ),
        (
            "SMK TELKOM JAKARTA",
            "West Jakarta, Indonesia",
            "Vocational High School Diploma in Software Engineering (Rekayasa Perangkat Lunak — RPL)",
            "July 2019 – June 2022"
        )
    ]

    for inst, loc, degree, period in edu_items:
        ep = doc.add_paragraph()
        ep.paragraph_format.space_before = Pt(4)
        ep.paragraph_format.space_after = Pt(1)
        ep.paragraph_format.keep_with_next = True
        
        inst_run = ep.add_run(inst)
        inst_run.font.bold = True
        inst_run.font.size = Pt(10)
        inst_run.font.color.rgb = PRIMARY_COLOR

        loc_run = ep.add_run(f" — {loc}")
        loc_run.font.size = Pt(9.5)
        loc_run.font.color.rgb = MUTED_TEXT

        p_run = ep.add_run(f" \t{period}")
        p_run.font.bold = True
        p_run.font.size = Pt(9)
        p_run.font.color.rgb = SECONDARY_COLOR

        dp = doc.add_paragraph()
        dp.paragraph_format.space_before = Pt(0)
        dp.paragraph_format.space_after = Pt(4)
        dp.paragraph_format.left_indent = Inches(0.15)
        deg_run = dp.add_run(degree)
        deg_run.font.size = Pt(9.5)
        deg_run.font.color.rgb = DARK_TEXT

    # Save document
    filename = "Hamza_Deleon_Wiradarma_Resume.docx"
    doc.save(filename)
    print(f"Resume generated successfully as '{filename}'")

if __name__ == '__main__':
    build_resume()
