import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set inner cell padding in dxa (1 pt = 20 dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for edge, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{edge}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_background(cell, fill_hex):
    """Set cell background color."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def remove_table_borders(table):
    """Remove all visible borders from a table."""
    tblPr = table._tbl.tblPr
    tblBorders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>\n'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>\n'
        f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>\n'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>\n'
        f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>\n'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>\n'
        f'</w:tblBorders>'
    )
    tblPr.append(tblBorders)

def add_bottom_border(paragraph, color_hex="2563EB", size_pt=1.5):
    """Add bottom accent line under section titles."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(int(size_pt * 8)))
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)

def build_professional_resume():
    doc = Document()

    # 0.5 in margins for sleek modern layout
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)

    # Color Palette: Executive Tech Theme
    NAVY = RGBColor(15, 23, 42)        # #0F172A Deep Navy / Slate Title
    BLUE = RGBColor(37, 99, 235)       # #2563EB Slate Accent Blue
    CHARCOAL = RGBColor(51, 65, 85)    # #334155 Dark Slate Text
    MUTED = RGBColor(100, 116, 139)    # #64748B Muted Slate
    BG_LIGHT_HEX = "F8FAFC"            # Light slate background fill
    BORDER_HEX = "E2E8F0"              # Subtle gray border

    FONT = 'Segoe UI'

    # Set base style font
    style = doc.styles['Normal']
    style.font.name = FONT
    style.font.size = Pt(9.5)
    style.font.color.rgb = CHARCOAL

    # -------------------------------------------------------------
    # HEADER BANNER (Table with subtle background card)
    # -------------------------------------------------------------
    header_tbl = doc.add_table(rows=1, cols=1)
    header_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_tbl.autofit = False
    
    cell = header_tbl.cell(0, 0)
    cell.width = Inches(7.4)
    set_cell_background(cell, "F1F5F9")  # Soft Slate Gray banner
    set_cell_margins(cell, top=180, bottom=180, left=240, right=240)

    # Rounded border styling box around header
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="single" w:sz="6" w:space="0" w:color="CBD5E1"/>\n'
        f'  <w:left w:val="single" w:sz="18" w:space="0" w:color="2563EB"/>\n'  # Blue left accent line
        f'  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="CBD5E1"/>\n'
        f'  <w:right w:val="single" w:sz="6" w:space="0" w:color="CBD5E1"/>\n'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)

    hp = cell.paragraphs[0]
    hp.paragraph_format.space_before = Pt(0)
    hp.paragraph_format.space_after = Pt(2)
    
    r_name = hp.add_run("HAMZA DELEON WIRADARMA")
    r_name.font.name = FONT
    r_name.font.size = Pt(20)
    r_name.font.bold = True
    r_name.font.color.rgb = NAVY

    sub_p = cell.add_paragraph()
    sub_p.paragraph_format.space_before = Pt(0)
    sub_p.paragraph_format.space_after = Pt(8)
    r_sub = sub_p.add_run("FULL STACK DEVELOPER")
    r_sub.font.name = FONT
    r_sub.font.size = Pt(11)
    r_sub.font.bold = True
    r_sub.font.color.rgb = BLUE

    # Contact items row
    meta_p = cell.add_paragraph()
    meta_p.paragraph_format.space_before = Pt(0)
    meta_p.paragraph_format.space_after = Pt(0)
    
    contact_items = [
        ("📍 Tangerang, Banten, 15116"),
        ("📱 +62 831-9463-4063"),
        ("🌐 Indonesian"),
        ("🎂 Feb 23, 2006")
    ]
    
    for idx, item in enumerate(contact_items):
        r = meta_p.add_run(item)
        r.font.name = FONT
        r.font.size = Pt(9)
        r.font.color.rgb = CHARCOAL
        
        if idx < len(contact_items) - 1:
            sep = meta_p.add_run("  •  ")
            sep.font.name = FONT
            sep.font.size = Pt(9)
            sep.font.color.rgb = BLUE
            sep.font.bold = True

    # Spacer
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after = Pt(4)

    # -------------------------------------------------------------
    # SECTION HEADING GENERATOR
    # -------------------------------------------------------------
    def add_section_header(title_text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        
        run = p.add_run(title_text.upper())
        run.font.name = FONT
        run.font.size = Pt(10.5)
        run.font.bold = True
        run.font.color.rgb = NAVY
        add_bottom_border(p, color_hex="2563EB", size_pt=1.2)
        return p

    # -------------------------------------------------------------
    # 1. PROFESSIONAL SUMMARY
    # -------------------------------------------------------------
    add_section_header("Professional Summary")
    
    p_sum = doc.add_paragraph()
    p_sum.paragraph_format.space_before = Pt(3)
    p_sum.paragraph_format.space_after = Pt(6)
    p_sum.paragraph_format.line_spacing = 1.15
    
    r_sum = p_sum.add_run(
        "Aspiring Full Stack Developer with a solid technical foundation in React.js, Express.js, Node.js, Python, PHP (Laravel), Java (Spring Boot), and relational/NoSQL databases. Proven track record of building performant web applications, REST APIs, interactive client-side systems, and decision-support algorithms through academic and freelance projects. Passionate about solving complex real-world problems, mastering modern software architectures, and writing clean, maintainable code. Seeking to leverage skills in a high-growth remote software engineering team while expanding expertise in full-stack architecture and AI systems."
    )
    r_sum.font.name = FONT
    r_sum.font.size = Pt(9.5)
    r_sum.font.color.rgb = CHARCOAL

    # -------------------------------------------------------------
    # 2. TECHNICAL SKILLS GRID (Formatted Table)
    # -------------------------------------------------------------
    add_section_header("Technical Skills")

    skills = [
        ("Frontend Development", "React.js (★★★★☆), JavaScript (ES6+), HTML5, CSS3 (Flexbox/Grid), TailwindCSS, Bootstrap"),
        ("Backend & Frameworks", "Express.js (★★★★☆), Python (★★★★☆), PHP / Laravel (★★★☆☆), Java / Spring Boot (★★★☆☆)"),
        ("Databases & Storage", "MySQL (★★★★☆), MongoDB, IndexedDB"),
        ("Tools & Methodologies", "Git & GitHub (★★★★☆), RESTful APIs, Jest, React Testing Library, Postman, Vite, Visual Basic")
    ]

    sk_tbl = doc.add_table(rows=len(skills), cols=2)
    sk_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(sk_tbl)
    
    for row_idx, (cat, val) in enumerate(skills):
        row = sk_tbl.rows[row_idx]
        
        # Category Column
        c1 = row.cells[0]
        c1.width = Inches(2.0)
        set_cell_margins(c1, top=60, bottom=60, left=60, right=60)
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_before = Pt(0)
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(f"•  {cat}")
        r1.font.name = FONT
        r1.font.bold = True
        r1.font.size = Pt(9)
        r1.font.color.rgb = BLUE

        # Values Column
        c2 = row.cells[1]
        c2.width = Inches(5.4)
        set_cell_margins(c2, top=60, bottom=60, left=60, right=60)
        p2 = c2.paragraphs[0]
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(val)
        r2.font.name = FONT
        r2.font.size = Pt(9)
        r2.font.color.rgb = CHARCOAL

    # -------------------------------------------------------------
    # 3. WORK EXPERIENCE
    # -------------------------------------------------------------
    add_section_header("Work Experience")

    # Helper function for perfect 2-column aligned job titles & dates
    def add_job_header(role_title, company_info, dates_str, location_str):
        tbl = doc.add_table(rows=1, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        remove_table_borders(tbl)
        
        c_left = tbl.cell(0, 0)
        c_left.width = Inches(5.2)
        set_cell_margins(c_left, top=80, bottom=40, left=0, right=60)
        
        p_left = c_left.paragraphs[0]
        p_left.paragraph_format.space_before = Pt(4)
        p_left.paragraph_format.space_after = Pt(0)
        p_left.paragraph_format.keep_with_next = True
        
        r1 = p_left.add_run(role_title)
        r1.font.name = FONT
        r1.font.bold = True
        r1.font.size = Pt(10.5)
        r1.font.color.rgb = NAVY
        
        r2 = p_left.add_run(f"  |  {company_info}")
        r2.font.name = FONT
        r2.font.size = Pt(9.5)
        r2.font.color.rgb = MUTED

        c_right = tbl.cell(0, 0)
        c_right = tbl.cell(0, 1)
        c_right.width = Inches(2.2)
        set_cell_margins(c_right, top=80, bottom=40, left=60, right=0)
        
        p_right = c_right.paragraphs[0]
        p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_right.paragraph_format.space_before = Pt(4)
        p_right.paragraph_format.space_after = Pt(0)
        p_right.paragraph_format.keep_with_next = True
        
        r3 = p_right.add_run(dates_str)
        r3.font.name = FONT
        r3.font.bold = True
        r3.font.size = Pt(9.5)
        r3.font.color.rgb = BLUE
        
        if location_str:
            p_right.add_run("\n")
            r4 = p_right.add_run(location_str)
            r4.font.name = FONT
            r4.font.size = Pt(8.5)
            r4.font.color.rgb = MUTED

    # --- JOB 1: Freelance Fullstack Developer ---
    add_job_header(
        role_title="Fullstack Developer",
        company_info="Freelance / Self-Employed",
        dates_str="2023 – Present",
        location_str="Tangerang, Banten"
    )

    freelance_projects = [
        (
            "Photon — Digital Image Processing Studio",
            "Designed and built a browser-based, desktop-class image processing platform combining a pixel-level HTML5 Canvas engine with a Python/Flask/OpenCV backend. Implemented 6 from-scratch edge detection algorithms, 5 custom binary image compression codecs (Huffman, LZW, RLE), YOLOv4-tiny AI object recognition, real-time RGB histogram analysis, and a framework-free Pub/Sub state management system supporting IndexedDB and MySQL storage."
        ),
        (
            "Decision Support System for Bandsaw Machine Procurement (MOORA Algorithm)",
            "Developed a full-stack web decision support application utilizing the MOORA mathematical decision algorithm to optimize industrial machinery purchasing. Built with PHP and MySQL, the system normalizes multi-attribute cost/benefit criteria (price, cutting capacity, power efficiency, durability) to generate objective, data-driven machine procurement rankings."
        ),
        (
            "Ogani — Full-Stack Organic E-Commerce Platform",
            "Built a complete full-stack e-commerce application specializing in organic grocery retail. Engineered customer-facing features including product catalog filtering, dynamic shopping cart, multi-step checkout, and order history tracking, paired with a comprehensive admin dashboard for real-time inventory management, order processing, and user authorization."
        )
    ]

    for p_name, p_desc in freelance_projects:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_before = Pt(2)
        bp.paragraph_format.space_after = Pt(3)
        bp.paragraph_format.left_indent = Inches(0.2)
        bp.paragraph_format.line_spacing = 1.12
        
        r_title = bp.add_run(f"{p_name}: ")
        r_title.font.name = FONT
        r_title.font.bold = True
        r_title.font.size = Pt(9)
        r_title.font.color.rgb = NAVY

        r_body = bp.add_run(p_desc)
        r_body.font.name = FONT
        r_body.font.size = Pt(9)
        r_body.font.color.rgb = CHARCOAL

    # --- JOB 2: Frontend Developer Intern ---
    add_job_header(
        role_title="Frontend Developer Intern",
        company_info="Meruya",
        dates_str="July 2021 – Nov 2021",
        location_str="Jakarta Barat, Indonesia"
    )

    intern_points = [
        "Developed responsive, accessible UI components using HTML5, CSS3 (Flexbox/Grid), and modern JavaScript (ES6+); built reusable React components that accelerated feature delivery by ~20%.",
        "Integrated front-end interfaces with RESTful APIs and implemented robust client-side validation and error handling, enhancing user task completion and reducing support tickets.",
        "Optimized front-end performance via lazy loading, code-splitting, and asset minification, reducing initial page load time by ~25%.",
        "Wrote unit and integration tests (Jest, React Testing Library) and maintained component documentation to improve codebase maintainability and team onboarding.",
        "Resolved cross-browser compatibility issues and addressed web accessibility (WCAG compliance) requirements prior to release."
    ]

    for pt in intern_points:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_before = Pt(2)
        bp.paragraph_format.space_after = Pt(2)
        bp.paragraph_format.left_indent = Inches(0.2)
        bp.paragraph_format.line_spacing = 1.12
        
        r_pt = bp.add_run(pt)
        r_pt.font.name = FONT
        r_pt.font.size = Pt(9)
        r_pt.font.color.rgb = CHARCOAL

    # -------------------------------------------------------------
    # 4. EDUCATION
    # -------------------------------------------------------------
    add_section_header("Education")

    education_list = [
        (
            "JAKARTA STATE POLYTECHNIC (Politeknik Negeri Jakarta - PNJ)",
            "Depok, Indonesia",
            "Bachelor of Applied Science in Informatics Engineering (Teknik Informatika)",
            "July 2022 – Present (Active)"
        ),
        (
            "CCIT — FACULTY OF ENGINEERING, UNIVERSITY OF INDONESIA (UI)",
            "Depok, Indonesia",
            "Professional Program in Information Technology",
            "July 2022 – July 2026"
        ),
        (
            "SMK TELKOM JAKARTA",
            "West Jakarta, Indonesia",
            "Vocational High School Diploma in Software Engineering (Rekayasa Perangkat Lunak)",
            "July 2019 – June 2022"
        )
    ]

    for inst, loc, deg, dates in education_list:
        tbl = doc.add_table(rows=1, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        remove_table_borders(tbl)
        
        cL = tbl.cell(0, 0)
        cL.width = Inches(5.2)
        set_cell_margins(cL, top=40, bottom=20, left=0, right=60)
        
        pL = cL.paragraphs[0]
        pL.paragraph_format.space_before = Pt(2)
        pL.paragraph_format.space_after = Pt(0)
        pL.paragraph_format.keep_with_next = True
        
        rI = pL.add_run(inst)
        rI.font.name = FONT
        rI.font.bold = True
        rI.font.size = Pt(9.5)
        rI.font.color.rgb = NAVY

        pD = cL.add_paragraph()
        pD.paragraph_format.space_before = Pt(0)
        pD.paragraph_format.space_after = Pt(3)
        rDg = pD.add_run(deg)
        rDg.font.name = FONT
        rDg.font.size = Pt(9)
        rDg.font.color.rgb = CHARCOAL

        cR = tbl.cell(0, 1)
        cR.width = Inches(2.2)
        set_cell_margins(cR, top=40, bottom=20, left=60, right=0)
        
        pR = cR.paragraphs[0]
        pR.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pR.paragraph_format.space_before = Pt(2)
        pR.paragraph_format.space_after = Pt(0)
        pR.paragraph_format.keep_with_next = True
        
        rDt = pR.add_run(dates)
        rDt.font.name = FONT
        rDt.font.bold = True
        rDt.font.size = Pt(9)
        rDt.font.color.rgb = BLUE
        
        pR.add_run("\n")
        rLc = pR.add_run(loc)
        rLc.font.name = FONT
        rLc.font.size = Pt(8.5)
        rLc.font.color.rgb = MUTED

    filename = "Hamza_Deleon_Wiradarma_Resume_Professional.docx"
    try:
        doc.save(filename)
        print(f"Professional resume generated: {filename}")
    except PermissionError:
        alt_filename = "Hamza_Deleon_Wiradarma_Resume_v2.docx"
        doc.save(alt_filename)
        print(f"Primary file was locked. Professional resume saved as: {alt_filename}")

if __name__ == '__main__':
    build_professional_resume()
